"""
================================================================================
👨‍🏫 智能教师作业批改助手 — AI Teacher Grading Assistant
================================================================================
一款基于 Streamlit + 多模型免费 API 的作业批改产品
支持：文本粘贴 / 文件上传(.txt .docx .pdf) / 批量批改 / 历史记录 / 结果导出

安装依赖：
    pip install streamlit openai python-docx PyPDF2

启动：
    streamlit run app.py

部署到 Streamlit Cloud（免费）：
    1. 将项目推送到 GitHub 仓库
    2. 访问 https://share.streamlit.io 连接仓库
    3. 自动获得公开链接 → 你的 AI Agent 就上线了！
================================================================================
"""

import streamlit as st
from openai import OpenAI
from docx import Document
import io
import datetime
import re
import base64
from pathlib import Path

# ============================================================================
# 页面配置
# ============================================================================
st.set_page_config(
    page_title="智能教师作业批改助手",
    page_icon="👨‍🏫",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============================================================================
# 自定义 CSS — 产品级 UI 美化
# ============================================================================
st.markdown(
    """
<style>
    /* ----- 全局字体和变量 ----- */
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@300;400;500;700;900&display=swap');

    html, body, [class*="css"] {
        font-family: 'Noto Sans SC', -apple-system, BlinkMacSystemFont, sans-serif;
    }

    /* ----- 主标题 ----- */
    .main-title {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: 2.6rem;
        font-weight: 900;
        text-align: center;
        margin-bottom: 0.3rem;
    }
    .sub-title {
        text-align: center;
        color: #888;
        font-size: 1rem;
        margin-bottom: 1.5rem;
    }

    /* ----- 卡片容器 ----- */
    .card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 16px;
        padding: 24px;
        margin: 12px 0;
        box-shadow: 0 4px 16px rgba(0,0,0,0.06);
    }

    /* ----- 统计数字 ----- */
    .stat-box {
        text-align: center;
        padding: 16px;
        border-radius: 12px;
        background: white;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    }
    .stat-number {
        font-size: 2rem;
        font-weight: 900;
        color: #667eea;
    }
    .stat-label {
        font-size: 0.85rem;
        color: #999;
    }

    /* ----- 按钮美化 ----- */
    .stButton > button {
        border-radius: 12px !important;
        font-weight: 700 !important;
        font-size: 1.1rem !important;
        padding: 12px 32px !important;
        transition: all 0.3s ease !important;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(102, 126, 234, 0.4);
    }

    /* ----- 结果区域 ----- */
    .result-container {
        background: white;
        border-radius: 16px;
        padding: 32px;
        margin-top: 20px;
        border: 1px solid #e0e0e0;
        box-shadow: 0 4px 20px rgba(0,0,0,0.05);
    }

    /* ----- 页脚 ----- */
    .footer {
        text-align: center;
        color: #bbb;
        font-size: 0.8rem;
        margin-top: 40px;
        padding: 20px 0;
        border-top: 1px solid #eee;
    }

    /* ----- 成功/警告 toast 自定义 ----- */
    div[data-testid="stNotification"] {
        border-radius: 12px !important;
    }
</style>
""",
    unsafe_allow_html=True,
)

# ============================================================================
# 初始化 Session State
# ============================================================================
DEFAULT_STATE = {
    "history": [],          # 批改历史记录
    "total_count": 0,       # 累计批改次数
    "last_result": "",      # 最近一次批改结果
    "last_subject": "",     # 最近一次批改学科
    "last_model": "",       # 最近一次使用的模型
    "last_homework": "",    # 最近一次批改的作业内容
}

for key, val in DEFAULT_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ============================================================================
# 模型配置映射
# ============================================================================
MODEL_CONFIG = {
    "DeepSeek-V3（硅基流动）": {
        "model": "deepseek-ai/DeepSeek-V3",
        "base_url": "https://api.siliconflow.cn/v1",
        "provider": "硅基流动 SiliconFlow",
    },
    "DeepSeek-R1（硅基流动）": {
        "model": "deepseek-ai/DeepSeek-R1",
        "base_url": "https://api.siliconflow.cn/v1",
        "provider": "硅基流动 SiliconFlow",
    },
    "Qwen2.5-72B（硅基流动）": {
        "model": "Qwen/Qwen2.5-72B-Instruct",
        "base_url": "https://api.siliconflow.cn/v1",
        "provider": "硅基流动 SiliconFlow",
    },
    "DeepSeek-V3（官方）": {
        "model": "deepseek-chat",
        "base_url": "https://api.deepseek.com",
        "provider": "DeepSeek 官方",
    },
    "Qwen-Max（阿里百炼）": {
        "model": "qwen-max-latest",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "provider": "阿里百炼",
    },
    "GLM-4（智谱AI）": {
        "model": "glm-4-plus",
        "base_url": "https://open.bigmodel.cn/api/paas/v4",
        "provider": "智谱AI ZhipuAI",
    },
    "Kimi（Moonshot）": {
        "model": "moonshot-v1-8k",
        "base_url": "https://api.moonshot.cn/v1",
        "provider": "Moonshot",
    },
}

# ============================================================================
# 学科 & 学段 — 系统提示词
# ============================================================================

GRADE_LEVELS = ["小学", "初中", "高中"]

SUBJECT_SYSTEM_PROMPTS = {
    "语文": {
        "icon": "📖",
        "prompt": """你是一位资深的{grade}语文教师，拥有 20 年教学经验。请对以下学生的语文作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **内容理解**（30分）：是否准确把握文章主旨，理解是否深入
2. **语言表达**（25分）：用词是否准确生动，句式是否丰富多变，修辞运用是否恰当
3. **结构逻辑**（20分）：段落层次是否清晰，论证是否严密，过渡是否自然
4. **书写规范**（15分）：错别字、标点符号使用是否正确，卷面是否整洁
5. **创意与亮点**（10分）：是否有独特的见解或优美的表达

请严格按以下格式输出，确保 Markdown 表格能正确渲染：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 内容理解 | 30 | XX | 一句话点评 |
| 语言表达 | 25 | XX | 一句话点评 |
| 结构逻辑 | 20 | XX | 一句话点评 |
| 书写规范 | 15 | XX | 一句话点评 |
| 创意亮点 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 📝 逐维度详细评语

> 以下对每个评分维度展开详细分析。

### 内容理解（XX/30）
（分析文章主旨把握、理解深度方面的表现，引用原文关键句作为佐证）

### 语言表达（XX/25）
（分析用词、句式、修辞手法方面的表现）

### 结构逻辑（XX/20）
（分析段落层次、论证逻辑方面的表现）

### 书写规范（XX/15）
（指出错别字、标点等问题，如有）

### 创意亮点（XX/10）
（评价独特见解和优美表达）

---

## ✨ 亮点

-
-

## ⚠️ 不足

-
-

---

## 💡 改进建议

1.
2.

---

## 🎯 教师寄语

> （一段温暖鼓励的总结寄语，针对{grade}学生的语气）""",
    },
    "数学": {
        "icon": "🔢",
        "prompt": """你是一位资深的{grade}数学教师，拥有 20 年教学经验。请对以下学生的数学作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **解题思路**（30分）：解题方法是否合理，逻辑是否清晰，是否选择了最优解法
2. **计算准确性**（30分）：计算过程是否正确，结果是否无误，单位是否统一
3. **步骤完整性**（20分）：解题步骤是否完整规范，有无跳步或遗漏
4. **规范性**（10分）：书写格式、数学符号使用是否规范
5. **拓展思维**（10分）：是否展现了多种解法或深入思考，能否举一反三

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 解题思路 | 30 | XX | 一句话点评 |
| 计算准确性 | 30 | XX | 一句话点评 |
| 步骤完整性 | 20 | XX | 一句话点评 |
| 规范性 | 10 | XX | 一句话点评 |
| 拓展思维 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析

（逐题给出「正确 ✅」或「有误 ❌」标记，分析解题过程，错误的给出正确解法）

---

## ✨ 亮点

-
-

## ⚠️ 不足

-
-

---

## 💡 改进建议

1.
2.

---

## 🎯 教师寄语

> （一段温暖鼓励的总结寄语）""",
    },
    "英语": {
        "icon": "🌍",
        "prompt": """你是一位资深的{grade}英语教师，拥有 20 年教学经验。请对以下学生的英语作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **语法准确性**（30分）：时态、语态、主谓一致、句型结构是否正确
2. **词汇运用**（25分）：词汇是否丰富恰当，搭配是否地道，拼写是否正确
3. **内容完整性**（20分）：是否覆盖题目要求，内容是否充实，逻辑是否合理
4. **篇章连贯**（15分）：段落衔接是否自然，过渡词使用是否恰当
5. **书写规范**（10分）：大小写、标点、格式是否正确

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 语法准确性 | 30 | XX | 一句话点评 |
| 词汇运用 | 25 | XX | 一句话点评 |
| 内容完整性 | 20 | XX | 一句话点评 |
| 篇章连贯 | 15 | XX | 一句话点评 |
| 书写规范 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 📝 逐维度详细评语

> 以下对每个评分维度展开详细分析。

### 语法准确性（XX/30）
（分析时态、语态、句型，摘录原文例句进行点评）

### 词汇运用（XX/25）
（分析词汇丰富度、搭配地道性，列出精彩用词和不当用词）

### 内容完整性（XX/20）
（分析是否覆盖题目要求，内容是否充实）

### 篇章连贯（XX/15）
（分析段落衔接和过渡词使用）

### 书写规范（XX/10）
（指出拼写错误、大小写、标点问题）

---

## ✨ 亮点

-
-

## ⚠️ 不足

-
-

---

## 💡 改进建议

1.
2.

---

## 🎯 教师寄语

> （用英文写一句鼓励 + 中文写一段寄语）""",
    },
    "物理": {
        "icon": "⚡",
        "prompt": """你是一位资深的{grade}物理教师，拥有 20 年教学经验。请对以下学生的物理作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **物理模型建立**（25分）：能否正确识别物理情境并建立模型
2. **公式运用**（25分）：公式选择是否正确，适用条件是否满足
3. **计算准确性**（20分）：数值计算是否正确，单位换算是否准确
4. **分析推理**（20分）：物理过程分析是否清晰，推理是否严密
5. **规范表达**（10分）：解题格式、符号使用、图像绘制是否规范

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 物理模型建立 | 25 | XX | 一句话点评 |
| 公式运用 | 25 | XX | 一句话点评 |
| 计算准确性 | 20 | XX | 一句话点评 |
| 分析推理 | 20 | XX | 一句话点评 |
| 规范表达 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析

（逐题分析解题过程，指出正确之处和错误之处，给出正确解法）

---

## ✨ 亮点
## ⚠️ 不足
## 💡 改进建议
## 🎯 教师寄语""",
    },
    "化学": {
        "icon": "🧪",
        "prompt": """你是一位资深的{grade}化学教师，拥有 20 年教学经验。请对以下学生的化学作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **化学原理理解**（25分）：对化学反应原理、规律的理解程度
2. **方程式书写**（25分）：化学方程式是否配平正确，条件标注是否完整
3. **计算能力**（20分）：化学计算（物质的量、浓度等）是否准确
4. **实验思维**（20分）：实验设计、现象分析、结论推理能力
5. **规范表达**（10分）：化学用语、符号、单位使用是否规范

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 化学原理理解 | 25 | XX | 一句话点评 |
| 方程式书写 | 25 | XX | 一句话点评 |
| 计算能力 | 20 | XX | 一句话点评 |
| 实验思维 | 20 | XX | 一句话点评 |
| 规范表达 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析
## ✨ 亮点
## ⚠️ 不足
## 💡 改进建议
## 🎯 教师寄语""",
    },
    "生物": {
        "icon": "🧬",
        "prompt": """你是一位资深的{grade}生物教师，拥有 20 年教学经验。请对以下学生的生物作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **概念理解**（30分）：对生物学概念、原理的理解是否准确深入
2. **知识应用**（25分）：能否将知识灵活应用到具体问题中
3. **逻辑分析**（20分）：对生物过程、实验结果的逻辑分析能力
4. **图表解读**（15分）：对生物图表、数据的信息提取和分析能力
5. **规范表达**（10分）：生物学专业术语使用是否规范

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 概念理解 | 30 | XX | 一句话点评 |
| 知识应用 | 25 | XX | 一句话点评 |
| 逻辑分析 | 20 | XX | 一句话点评 |
| 图表解读 | 15 | XX | 一句话点评 |
| 规范表达 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析
## ✨ 亮点
## ⚠️ 不足
## 💡 改进建议
## 🎯 教师寄语""",
    },
    "历史": {
        "icon": "📜",
        "prompt": """你是一位资深的{grade}历史教师，拥有 20 年教学经验。请对以下学生的历史作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **史实准确性**（30分）：历史事实、时间、人物是否准确
2. **分析深度**（25分）：对历史事件的原因、影响分析是否深入
3. **逻辑论证**（20分）：论述是否有条理，论证是否充分
4. **史料运用**（15分）：能否恰当引用史料支撑观点
5. **表达规范**（10分）：语言表达是否清晰、学术规范

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 史实准确性 | 30 | XX | 一句话点评 |
| 分析深度 | 25 | XX | 一句话点评 |
| 逻辑论证 | 20 | XX | 一句话点评 |
| 史料运用 | 15 | XX | 一句话点评 |
| 表达规范 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析
## ✨ 亮点
## ⚠️ 不足
## 💡 改进建议
## 🎯 教师寄语""",
    },
    "地理": {
        "icon": "🌏",
        "prompt": """你是一位资深的{grade}地理教师，拥有 20 年教学经验。请对以下学生的地理作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **地理知识理解**（30分）：对地理概念、原理的理解是否准确
2. **空间思维**（25分）：空间定位、地理分布、区域分析能力
3. **综合分析**（20分）：多要素综合分析和因果关系推理能力
4. **图表技能**（15分）：地图、统计图表的解读和绘制能力
5. **表达规范**（10分）：地理术语使用是否规范

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 地理知识理解 | 30 | XX | 一句话点评 |
| 空间思维 | 25 | XX | 一句话点评 |
| 综合分析 | 20 | XX | 一句话点评 |
| 图表技能 | 15 | XX | 一句话点评 |
| 表达规范 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析
## ✨ 亮点
## ⚠️ 不足
## 💡 改进建议
## 🎯 教师寄语""",
    },
    "政治": {
        "icon": "⚖️",
        "prompt": """你是一位资深的{grade}政治教师，拥有 20 年教学经验。请对以下学生的政治作业进行多维度批改。

批改维度与评分标准（总分 100 分）：
1. **理论理解**（30分）：对政治理论、概念的准确把握程度
2. **联系实际**（25分）：能否用理论分析现实问题
3. **逻辑论证**（20分）：论证是否严谨、条理是否清晰
4. **辩证思维**（15分）：能否多角度、辩证地看待问题
5. **表达规范**（10分）：政治术语使用是否规范

请严格按以下格式输出：

---

## 📊 得分总览

| 批改维度 | 满分 | 得分 | 简要说明 |
| :------ | :--: | :--: | :------ |
| 理论理解 | 30 | XX | 一句话点评 |
| 联系实际 | 25 | XX | 一句话点评 |
| 逻辑论证 | 20 | XX | 一句话点评 |
| 辩证思维 | 15 | XX | 一句话点评 |
| 表达规范 | 10 | XX | 一句话点评 |
| **🎯 总分** | **100** | **XX** | |

---

## 🔍 逐题分析
## ✨ 亮点
## ⚠️ 不足
## 💡 改进建议
## 🎯 教师寄语""",
    },
}

# ============================================================================
# 工具函数
# ============================================================================


def parse_txt(uploaded_file) -> str:
    """解析 .txt 文件，自动处理 UTF-8 / GBK 编码。"""
    raw = uploaded_file.read()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("gbk", errors="ignore")


def parse_docx(uploaded_file) -> str:
    """解析 .docx 文件，提取所有段落和表格中的文字。"""
    doc = Document(uploaded_file)
    parts = []

    # 提取段落
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def parse_pdf(uploaded_file) -> str:
    """解析 .pdf 文件（尝试 PyPDF2，失败则提示安装）。"""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(uploaded_file)
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts)
    except ImportError:
        raise ImportError("需要安装 PyPDF2：pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败：{e}")


def extract_text_from_file(uploaded_file) -> str:
    """根据文件扩展名自动选择解析方式。"""
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".txt"):
        return parse_txt(uploaded_file)
    elif file_name.endswith(".docx"):
        return parse_docx(uploaded_file)
    elif file_name.endswith(".pdf"):
        return parse_pdf(uploaded_file)
    else:
        raise ValueError(f"不支持的文件格式「{file_name}」，仅支持 .txt / .docx / .pdf")


def detect_subject(content: str) -> str | None:
    """根据关键词简单推测作业内容的学科。返回学科名或 None。"""
    text = content[:500]

    # 学科特征关键词
    signals = {
        "数学": [
            "解：", "证明：", "方程", "函数", "求导", "积分", "极限", "矩阵", "向量",
            "sin", "cos", "tan", "log", "lim", "dx", "dy", "x=", "y=", "f(x)",
            "三角形", "圆", "面积", "体积", "概率", "统计", "数列", "不等式",
            "√", "∑", "∫", "π", "∞", "≤", "≥", "≠", "±", "×", "÷",
        ],
        "英语": [
            "happy", "school", "friend", "family", "travel", "beautiful",
            " I ", " you ", " he ", " she ", " we ", " they ", " the ",
            " is ", " are ", " was ", " were ", " have ", " has ", " do ", " does ",
            "what", "when", "where", "which", "because", "although", "however",
            "a ", "an ", "the ", "this ", "that ", "my ", "your ",
            "essay", "writing", "letter", "Dear ", "Sincerely",
        ],
        "物理": [
            "牛顿", "加速度", "力", "质量", "速度", "位移", "功", "能量", "动量",
            "电场", "磁场", "电流", "电压", "电阻", "欧姆", "安培", "伏特",
            "F=", "ma", "m/s", "m/s²", "W=", "J", "N", "Pa",
        ],
        "化学": [
            "化学方程式", "反应", "mol", "NaCl", "H₂O", "CO₂", "NaOH", "H₂SO₄",
            "氧化", "还原", "酸", "碱", "盐", "元素周期", "原子", "分子",
            "质量守恒", "浓度", "摩尔", "pH", "沉淀", "气体", "溶液",
        ],
        "生物": [
            "细胞", "DNA", "RNA", "基因", "蛋白质", "酶", "光合作用", "呼吸作用",
            "染色体", "遗传", "变异", "进化", "生态系统", "种群", "神经",
        ],
        "历史": [
            "朝代", "皇帝", "战争", "革命", "条约", "变法", "起义",
            "公元", "世纪", "年代", "古代", "近代", "清朝", "唐朝",
        ],
        "地理": [
            "气候", "地形", "经纬度", "海拔", "洋流", "板块", "地震",
            "人口", "城市化", "农业", "工业", "资源", "环境",
        ],
        "政治": [
            "国家", "政府", "法律", "权利", "义务", "民主", "制度",
            "社会主义", "市场经济", "宪法", "公民", "价值观",
        ],
    }

    scores = {}
    for subject, keywords in signals.items():
        count = sum(1 for kw in keywords if kw.lower() in text.lower())
        if count > 0:
            scores[subject] = count

    if not scores:
        return None

    best = max(scores, key=scores.get)

    # 只有信号足够强（>=3个关键词命中）才认为有效
    if scores[best] >= 3:
        return best
    return None


def call_llm(api_key: str, model_name: str, system_prompt: str, user_content: str) -> str:
    """使用 OpenAI 兼容接口调用大模型。"""
    config = MODEL_CONFIG[model_name]

    client = OpenAI(
        api_key=api_key,
        base_url=config["base_url"],
        timeout=120.0,
    )

    response = client.chat.completions.create(
        model=config["model"],
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ],
        temperature=0.7,
        max_tokens=4096,
    )

    return response.choices[0].message.content


def get_download_link(text: str, filename: str, label: str = "📥 下载结果") -> str:
    """生成 Markdown 文件下载链接。"""
    b64 = base64.b64encode(text.encode("utf-8")).decode()
    ext = filename.split(".")[-1]
    mime = "text/markdown" if ext == "md" else "text/plain"
    return f'<a href="data:{mime};base64,{b64}" download="{filename}" style="text-decoration:none;">{label}</a>'


# ============================================================================
# 侧边栏 — 系统配置
# ============================================================================
with st.sidebar:
    st.image(
        "https://img.icons8.com/fluency/96/teacher.png",
        width=72,
    )
    st.markdown("## ⚙️ 系统配置")

    # ---- 模型选择 ----
    st.markdown("### 🤖 选择大模型")

    model_names = list(MODEL_CONFIG.keys())

    model_index = st.selectbox(
        "选择大模型",
        options=range(len(model_names)),
        format_func=lambda i: model_names[i],
        index=0,
        help="选择用于批改作业的大语言模型",
        key="model_selector",
    )
    model_choice = model_names[model_index]
    model_cfg = MODEL_CONFIG[model_choice]

    st.caption(f"提供商：{model_cfg['provider']}")

    # ---- 学科选择 ----
    st.markdown("### 📚 批改设置")

    subject_choice = st.selectbox(
        "选择批改学科",
        options=list(SUBJECT_SYSTEM_PROMPTS.keys()),
        format_func=lambda s: f"{SUBJECT_SYSTEM_PROMPTS[s]['icon']} {s}",
        index=0,
        help="选择作业对应的学科",
        key="subject_selector",
    )

    grade_choice = st.selectbox(
        "选择学段",
        options=GRADE_LEVELS,
        index=1,
        help="选择学生所在学段，以适配批改难度和措辞",
        key="grade_selector",
    )

    # ---- API Key ----
    st.markdown("### 🔑 API Key")
    api_key = st.text_input(
        "请输入 API Key",
        type="password",
        placeholder="sk-xxxxxxxxxxxxxxxx",
        help=f"在 {model_cfg['provider']} 平台获取 API Key",
        key="api_key_input",
    )

    st.divider()

    # ---- 使用说明 ----
    with st.expander("📖 使用说明", expanded=True):
        st.caption("**4 步完成批改：**")
        st.caption("1️⃣ 选择大模型")
        st.caption("2️⃣ 选择学科和学段")
        st.caption("3️⃣ 输入 API Key")
        st.caption("4️⃣ 粘贴/上传作业 → 开始批改")

    with st.expander("📊 批改统计", expanded=False):
        st.metric("累计批改次数", st.session_state.total_count)
        if st.session_state.history:
            st.caption(f"最近批改：{st.session_state.history[-1]['time']}")

    # ---- 清除历史 ----
    if st.button("🗑️ 清除批改历史", use_container_width=True):
        st.session_state.history = []
        st.session_state.total_count = 0
        st.session_state.last_result = ""
        st.rerun()

# ============================================================================
# 主页面
# ============================================================================
st.markdown('<p class="main-title">👨‍🏫 智能教师作业批改助手</p>', unsafe_allow_html=True)
st.markdown(
    '<p class="sub-title">AI 驱动的多维度作业批改 · 支持 9 大学科 · 7 款大模型</p>',
    unsafe_allow_html=True,
)

# ----- 功能介绍卡片 -----
col_a, col_b, col_c, col_d = st.columns(4)
with col_a:
    st.markdown(
        '<div class="stat-box"><div class="stat-number">9+</div><div class="stat-label">支持学科</div></div>',
        unsafe_allow_html=True,
    )
with col_b:
    st.markdown(
        '<div class="stat-box"><div class="stat-number">7</div><div class="stat-label">可选模型</div></div>',
        unsafe_allow_html=True,
    )
with col_c:
    st.markdown(
        '<div class="stat-box"><div class="stat-number">3</div><div class="stat-label">输入方式</div></div>',
        unsafe_allow_html=True,
    )
with col_d:
    st.markdown(
        '<div class="stat-box"><div class="stat-number">5</div><div class="stat-label">评分维度</div></div>',
        unsafe_allow_html=True,
    )

# ----- Tab 切换输入方式 -----
tab1, tab2, tab3 = st.tabs(["📝 文本直接粘贴", "📂 单文件上传", "📚 批量上传（多份作业）"])

homework_content = ""

with tab1:
    text_input = st.text_area(
        "请在此粘贴学生作业内容",
        height=300,
        placeholder="请将学生的作业文字粘贴到此处...\n\n例如：\n· 语文：学生作文《我的家乡》\n· 数学：方程 x²+5x+6=0 的解题过程\n· 英语：My Favorite Season 写作",
        key="text_input",
    )

with tab2:
    uploaded_file = st.file_uploader(
        "上传作业文件（支持 .txt / .docx / .pdf）",
        type=["txt", "docx", "pdf"],
        help="拖拽文件到此处或点击浏览选择文件",
        key="single_file_uploader",
    )

    if uploaded_file is not None:
        file_detail = st.container()
        file_detail.info(
            f"✅ 已上传：**{uploaded_file.name}**（{uploaded_file.size / 1024:.1f} KB）"
        )
        # 预览前 500 字
        try:
            preview_text = extract_text_from_file(uploaded_file)
            with st.expander("👀 内容预览（前 500 字）", expanded=False):
                st.text(preview_text[:500] + ("..." if len(preview_text) > 500 else ""))
        except Exception:
            pass

with tab3:
    uploaded_files = st.file_uploader(
        "批量上传多份作业（支持 .txt / .docx / .pdf）",
        type=["txt", "docx", "pdf"],
        accept_multiple_files=True,
        help="可同时选择多个文件，逐一批改",
        key="batch_file_uploader",
    )

    if uploaded_files:
        st.info(f"✅ 已选择 **{len(uploaded_files)}** 份作业")
        file_names = [f.name for f in uploaded_files]
        st.caption("📋 文件列表：" + "、".join(file_names))

# ----- 提交按钮 -----
st.divider()
col_l, col_c, col_r = st.columns([1, 1, 1])
with col_c:
    submit_btn = st.button(
        "🚀 开始智能批改",
        type="primary",
        use_container_width=True,
    )

# ============================================================================
# 核心逻辑 — 批改
# ============================================================================
if submit_btn:
    # ---- 1. 校验 ----
    errors = []

    if not api_key.strip():
        errors.append("请先在侧边栏输入您的 **API Key**。")

    # 判断输入来源
    input_mode = None  # "text" | "single" | "batch"
    homework_list = []  # [(filename, content), ...]

    if text_input.strip():
        input_mode = "text"
        homework_list = [("📝 粘贴文本", text_input.strip())]
    elif uploaded_file is not None:
        input_mode = "single"
        try:
            content = extract_text_from_file(uploaded_file)
            if not content.strip():
                errors.append("文件内容为空，请检查文件。")
            homework_list = [(uploaded_file.name, content)]
        except ValueError as e:
            errors.append(str(e))
        except ImportError as e:
            errors.append(str(e))
        except Exception as e:
            errors.append(f"文件解析失败：{e}")
    elif uploaded_files:
        input_mode = "batch"
        for uf in uploaded_files:
            try:
                content = extract_text_from_file(uf)
                if content.strip():
                    homework_list.append((uf.name, content))
                else:
                    st.warning(f"⚠️ 文件「{uf.name}」内容为空，已跳过")
            except Exception as e:
                st.warning(f"⚠️ 文件「{uf.name}」解析失败：{e}")
        if not homework_list:
            errors.append("所有上传文件解析失败或内容为空。")
    else:
        errors.append("请在 Tab 中粘贴作业内容，或上传作业文件。")

    if errors:
        for err in errors:
            st.warning(f"⚠️ {err}")
        st.stop()

    # ---- 2. 学科内容匹配检测 ----
    # 取第一份作业内容做检测
    first_content = homework_list[0][1]
    detected = detect_subject(first_content)
    if detected and detected != subject_choice:
        st.warning(
            f"🤔 检测到作业内容更像是 **{detected}** 作业，"
            f"但你选择的是 **{subject_choice}**。\n\n"
            f"如果内容确实属于 {subject_choice}，请忽略此提示继续批改。"
        )

    # ---- 3. 获取系统提示词 ----
    system_prompt = SUBJECT_SYSTEM_PROMPTS[subject_choice]["prompt"].format(grade=grade_choice)

    # ---- 4. 逐份批改 ----
    all_results = []

    for idx, (file_name, content) in enumerate(homework_list):
        batch_label = f"（{idx+1}/{len(homework_list)}）" if len(homework_list) > 1 else ""
        status_text = f"🤔 正在用 **{model_choice}** 批改 {subject_choice} 作业「{file_name}」{batch_label}..."

        with st.spinner(status_text):
            try:
                result = call_llm(
                    api_key=api_key.strip(),
                    model_name=model_choice,
                    system_prompt=system_prompt,
                    user_content=f"【学科】{subject_choice}\n【学段】{grade_choice}\n【作业来源】{file_name}\n\n【学生作业内容】\n{content}",
                )
                all_results.append((file_name, content, result))
                st.session_state.total_count += 1

                # 保存历史
                st.session_state.history.append(
                    {
                        "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "subject": subject_choice,
                        "grade": grade_choice,
                        "model": model_choice,
                        "file_name": file_name,
                        "content": content[:200],
                        "result": result,
                    }
                )
                # 保持历史记录不超过 50 条
                if len(st.session_state.history) > 50:
                    st.session_state.history = st.session_state.history[-50:]

            except Exception as e:
                error_msg = str(e)
                st.error(f"❌ 批改「{file_name}」失败：{error_msg}")
                all_results.append((file_name, content, f"**批改失败**：{error_msg}"))

    # ---- 5. 展示结果 ----
    st.divider()

    for idx, (file_name, content, result) in enumerate(all_results):
        if len(all_results) > 1:
            st.subheader(f"📊 批改结果 {idx+1}：{file_name}")

        # 原始作业（折叠）
        with st.expander(f"📋 查看原始作业「{file_name}」", expanded=False):
            st.text(content[:2000] + ("\n\n...（内容过长，已截断）" if len(content) > 2000 else ""))

        # 批改结果
        with st.container(border=True):
            st.markdown(result)

        # 下载按钮
        download_text = f"# 教师作业批改报告\n\n"
        download_text += f"- **学科**：{subject_choice}\n"
        download_text += f"- **学段**：{grade_choice}\n"
        download_text += f"- **批改模型**：{model_choice}\n"
        download_text += f"- **批改时间**：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        download_text += f"- **作业来源**：{file_name}\n\n"
        download_text += "---\n\n"
        download_text += result

        safe_name = re.sub(r"[^\w\-.]", "_", file_name)
        st.markdown(
            get_download_link(
                download_text,
                f"批改报告_{safe_name}_{datetime.datetime.now().strftime('%Y%m%d%H%M')}.md",
                f"📥 下载批改报告（Markdown）",
            ),
            unsafe_allow_html=True,
        )

        if idx < len(all_results) - 1:
            st.divider()

    # ---- 6. 保存最近结果 ----
    if all_results:
        st.session_state.last_result = all_results[-1][2]
        st.session_state.last_subject = subject_choice
        st.session_state.last_model = model_choice
        st.session_state.last_homework = all_results[-1][1]

# ============================================================================
# 历史记录区（仅在无新批改时展示）
# ============================================================================
if not submit_btn and st.session_state.history:
    st.divider()
    st.subheader("📜 批改历史记录")

    history_df = [
        {
            "时间": h["time"],
            "学科": h["subject"],
            "学段": h["grade"],
            "模型": h["model"],
            "文件": h["file_name"],
            "内容预览": h["content"][:50] + "...",
        }
        for h in reversed(st.session_state.history[-10:])
    ]

    for i, h in enumerate(reversed(st.session_state.history[-10:])):
        with st.expander(f"{h['time']} | {h['subject']} | {h['file_name']}", expanded=False):
            st.caption(f"**模型**：{h['model']} | **学段**：{h['grade']}")
            st.markdown(h["result"])

# ============================================================================
# 页脚
# ============================================================================
st.markdown(
    """
<div class="footer">
    👨‍🏫 <strong>智能教师作业批改助手</strong> v2.0 &nbsp;|&nbsp;
    支持 9 大学科 &nbsp;|&nbsp;
    集成 7 款大模型 &nbsp;|&nbsp;
    Made with ❤️ + Streamlit + LLM
</div>
""",
    unsafe_allow_html=True,
)
